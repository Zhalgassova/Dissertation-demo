from io import BytesIO

from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill
from openpyxl.utils import get_column_letter


def build_template_workbook():
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = 'Experiment Template'

    headers = ['���', '���']
    for lesson in range(1, 9):
        for task in range(1, 4):
            headers.append(f'{lesson} ����? {task} ��������')

    sheet.append(headers)
    sheet.append(['������ ����?���', '�����������'] + [4, 5, 4] * 8)
    sheet.append(['�?��?���� ���?��?��', '��?����'] + [3, 4, 3] * 8)

    header_fill = PatternFill('solid', fgColor='186F65')
    for cell in sheet[1]:
        cell.font = Font(color='FFFFFF', bold=True)
        cell.fill = header_fill

    sheet.freeze_panes = 'A2'
    sheet.column_dimensions['A'].width = 28
    sheet.column_dimensions['B'].width = 18
    for index in range(3, len(headers) + 1):
        sheet.column_dimensions[get_column_letter(index)].width = 15

    output = BytesIO()
    workbook.save(output)
    output.seek(0)
    return output


def build_report_docx(dataset, summary):
    document = Document()
    normal_style = document.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(12)

    document.add_heading('����������? �?�� ������� �����������? ����', level=1)
    document.add_paragraph(f'������ �����: {dataset.title}')
    document.add_paragraph(f'����: {dataset.original_filename}')
    if dataset.sheet_name:
        document.add_paragraph(f'����?: {dataset.sheet_name}')

    metrics = summary.get('metrics', {})
    comparison = summary.get('comparison', {})
    prepost = summary.get('prepost', {})
    risk_zones = summary.get('risk_zones', {})
    recommendations = summary.get('recommendations', [])
    deficits = summary.get('topic_deficits', [])
    tables = summary.get('tables', {})

    document.add_heading('1. ����� �?���������', level=2)
    document.add_paragraph(
        f"�?������ ����: {metrics.get('students', 0)}. "
        f"��� ����: {metrics.get('groups', 0)}. "
        f"����? ����: {metrics.get('sessions', 0)}. "
        f"������ �?����: {metrics.get('average_score', 0)}%."
    )

    document.add_heading('2. ������������� ���������', level=2)
    document.add_paragraph(
        f"{comparison.get('basis', '-')}. "
        f"p-value = {comparison.get('p_value', '-')}, "
        f"effect size = {comparison.get('effect_size', '-')}. "
        f"{comparison.get('interpretation', '')}"
    )

    document.add_heading('3. Pre-test / Post-test ?����������', level=2)
    if prepost.get('group_pairs'):
        table = document.add_table(rows=1, cols=4)
        hdr = table.rows[0].cells
        hdr[0].text = '���'
        hdr[1].text = 'Pre-test'
        hdr[2].text = 'Post-test'
        hdr[3].text = '?��'
        for item in prepost['group_pairs']:
            row = table.add_row().cells
            row[0].text = str(item['group_name'])
            row[1].text = str(item['pre_value'])
            row[2].text = str(item['post_value'])
            row[3].text = str(item['gain'])
    else:
        document.add_paragraph('Pre-test/Post-test ?��� ������� ����� ���������.')

    document.add_heading('4. ?��� ����?����', level=2)
    document.add_paragraph(
        f"��?��� �?�����: {risk_zones.get('high', 0)}; "
        f"������ �?�����: {risk_zones.get('medium', 0)}; "
        f"�?��� �?�����: {risk_zones.get('low', 0)}."
    )

    document.add_heading('5. ?��� ��?�������', level=2)
    for item in deficits[:8]:
        document.add_paragraph(
            f"{item['group_name']} ����: {item['lesson_topic']} ({item['session_label']}) - "
            f"{item['avg_score']}%, ��?���: {item['status']}.",
            style='List Bullet',
        )

    document.add_heading('6. �?����� ��?��� �?������', level=2)
    for item in tables.get('weakest_students', [])[:8]:
        document.add_paragraph(
            f"{item['student_name']} ({item['group_name']}): ������ {item['avg_score']}%, "
            f"�?����� {item['risk_score']}%, ?��� ����? - {item.get('weakest_session', '-')}, "
            f"?��� �?�� - {item.get('weakest_topic', '-')}.",
            style='List Bullet',
        )

    document.add_heading('7. ������������? ?��������', level=2)
    for item in recommendations[:10]:
        document.add_paragraph(item, style='List Bullet')

    document.add_heading('8. ?��������', level=2)
    document.add_paragraph(
        '�?� �������� ��� ���� ��������� ������ �?������ ������������? ����� ?���������� '
        '��������������: ?��� �?��������, ���������� ����?����� �?�� ��?������? ����������� �?������, '
        '������-�? ����������������?�� ?�������� �����.'
    )

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output
